"""create_document(.docx) — 렌더 검증 + 생성-편집 왕복(roundtrip)."""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from document_adapter import create_document, load
from document_adapter.generate import docx_from_markdown

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

SAMPLE_MD = """# 주간 회의록

## 참석자
| 이름 | 부서 | 비고 |
|---|---|---|
| 김경윤 | AI플랫폼 | 주재 |
| 홍길동 | 개발1팀 | |

## 결정 사항
- 문서 생성 기능 **v0.15** 이번 주 배포
- HWPX 지원은 *Phase 2* 로 분리

## 액션 아이템
1. 작업계획서 리뷰

> 다음 회의: 7/15

---

```
uv build --wheel
```
"""


def test_docx_styles_and_structure() -> None:
    data = docx_from_markdown(SAMPLE_MD, lang="ko")
    doc = Document(io.BytesIO(data))

    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert styles[0] == "Heading 1"
    assert "Heading 2" in styles
    assert "List Bullet" in styles
    assert "List Number" in styles
    assert "Intense Quote" in styles

    # 표: 헤더행 굵게
    assert len(doc.tables) == 1
    header_cell = doc.tables[0].cell(0, 0)
    assert header_cell.text == "이름"
    assert all(r.bold for r in header_cell.paragraphs[0].runs)
    assert doc.tables[0].cell(1, 0).text == "김경윤"

    # 인라인 서식이 run 으로 살아있는지
    bullet_runs = [
        r for p in doc.paragraphs if p.style.name == "List Bullet" for r in p.runs
    ]
    assert any(r.text == "v0.15" and r.bold for r in bullet_runs)
    assert any(r.text == "Phase 2" and r.italic for r in bullet_runs)


def test_korean_base_font_with_east_asia() -> None:
    doc = Document(io.BytesIO(docx_from_markdown("# 제목\n본문", lang="ko")))
    normal = doc.styles["Normal"]
    assert normal.font.name == "맑은 고딕"
    assert normal.element.rPr.rFonts.get(f"{{{_W_NS}}}eastAsia") == "맑은 고딕"


def test_english_base_font() -> None:
    doc = Document(io.BytesIO(docx_from_markdown("# Title\nBody", lang="en-US")))
    assert doc.styles["Normal"].font.name == "Calibri"


def test_hr_renders_bottom_border() -> None:
    doc = Document(io.BytesIO(docx_from_markdown("위\n\n---\n\n아래", lang="ko")))
    borders = [
        p for p in doc.paragraphs
        if p._p.pPr is not None and p._p.pPr.find(f"{{{_W_NS}}}pBdr") is not None
    ]
    assert len(borders) == 1


def test_empty_markdown_raises_bilingual() -> None:
    with pytest.raises(ValueError) as exc:
        docx_from_markdown("   \n  ", lang="ko")
    assert "비어" in str(exc.value)


def test_roundtrip_generate_then_edit(tmp_path: Path) -> None:
    """원칙 3: 생성 직후 load() 좌표 편집이 성립해야 한다."""
    out = create_document(tmp_path / "회의록.docx", markdown=SAMPLE_MD, lang="ko")
    assert out.exists()

    adapter = load(out)
    try:
        tables = adapter.get_tables(preview_rows=10)
        assert len(tables) == 1
        assert tables[0].preview[1][0] == "김경윤"

        # 셀 수정 + 행 삽입 (v0.14 insert_row) 이 생성 문서에서 동작
        adapter.set_cell(0, 2, 2, "참관")
        adapter.insert_row(0, ["유지수", "기획", ""], at_row=3)
        adapter.save()
    finally:
        adapter.close()

    verify = load(out)
    try:
        t = verify.get_tables(preview_rows=10)[0]
    finally:
        verify.close()
    assert t.rows == 4
    assert t.preview[2][2] == "참관"
    assert t.preview[3][0] == "유지수"
