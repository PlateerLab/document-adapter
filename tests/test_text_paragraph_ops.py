"""find_text / replace_text / insert_text / get_text_map 어댑터 테스트.

DOCX/HWPX 에서 표 밖 본문·표 셀·머리말의 임의 텍스트를 서식 보존으로
치환/삽입하는 v0.13 기능 검증. 픽스처는 그때그때 생성 (외부 리소스 없음).
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from hwpx.document import HwpxDocument

from document_adapter import load
from document_adapter.base import NotImplementedForFormat
from document_adapter.hwpx_core import HP_RUN, HP_T


# ==================== DOCX 픽스처 ====================

def _make_docx_body(path: Path) -> None:
    """제목 2개 + run 분할 문단 + 표 + 본문 문단으로 구성된 문서.

    '홍길' + '동님께' run 분할은 Word 편집 이력으로 단어가 쪼개진 상황 재현.
    """
    doc = Document()
    doc.add_heading("1. 신청인 정보", level=1)
    p = doc.add_paragraph()
    r1 = p.add_run("홍길")
    r1.bold = True
    p.add_run("동님께 안내드립니다")

    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "담당자"
    t.cell(0, 1).text = "내용"

    doc.add_heading("4. 결재란", level=1)
    doc.add_paragraph("담당자      팀장      부장")
    doc.save(path)


# ==================== DOCX: find_text ====================

def test_docx_find_text_run_split(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        matches = doc.find_text("홍길동")
    finally:
        doc.close()
    # run 분할('홍길'+'동님께')에도 매치
    assert len(matches) == 1
    m = matches[0]
    assert m.scope == "body"
    assert "«홍길동»" in m.context
    assert m.nearest_heading == "1. 신청인 정보"


def test_docx_find_text_multiple_with_heading_context(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        matches = doc.find_text("담당자")
    finally:
        doc.close()
    assert len(matches) == 2
    # 1건: 표 셀 (신청인 정보 섹션), 1건: 본문 (결재란 섹션)
    assert matches[0].scope == "table"
    assert matches[0].location.startswith("table[0].cell(0,0)")
    assert matches[0].nearest_heading == "1. 신청인 정보"
    assert matches[1].scope == "body"
    assert matches[1].nearest_heading == "4. 결재란"
    # match_index 는 등장 순번
    assert [m.match_index for m in matches] == [0, 1]


def test_docx_find_text_scope_filter(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        only_table = doc.find_text("담당자", scope="table")
        only_body = doc.find_text("담당자", scope="body")
    finally:
        doc.close()
    assert len(only_table) == 1 and only_table[0].scope == "table"
    assert len(only_body) == 1 and only_body[0].scope == "body"


# ==================== DOCX: replace_text ====================

def test_docx_replace_preserves_run_format(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)

    doc = load(src)
    try:
        report = doc.replace_text("홍길동", "유지수")
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 1
    assert report["changes"][0]["matched"] == "홍길동"
    assert "유지수님께" in report["changes"][0]["paragraph_after"]

    # 재로드 후 run 구조/서식 검증
    d = Document(str(src))
    para = d.paragraphs[1]
    assert para.text == "유지수님께 안내드립니다"
    # 첫 run: 치환문이 들어가고 bold 유지
    assert para.runs[0].text == "유지수"
    assert para.runs[0].bold is True
    # 둘째 run: 매치 이후 부분만 남고 서식 불변(bold 아님)
    assert para.runs[1].text == "님께 안내드립니다"
    assert not para.runs[1].bold


def test_docx_replace_all_occurrences(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        report = doc.replace_text("담당자", "책임자")
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 2

    doc2 = load(src)
    try:
        assert doc2.find_text("담당자") == []
        assert len(doc2.find_text("책임자")) == 2
        # 표 셀도 치환됐는지
        assert doc2.get_cell(0, 0, 0).text == "책임자"
    finally:
        doc2.close()


def test_docx_replace_specific_occurrence(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        # 결재란 쪽(두 번째 등장)만 치환
        report = doc.replace_text("담당자", "책임자", occurrences=[1])
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 1
    assert report["changes"][0]["match_index"] == 1
    assert report["changes"][0]["scope"] == "body"

    doc2 = load(src)
    try:
        remain = doc2.find_text("담당자")
        assert len(remain) == 1
        assert remain[0].scope == "table"   # 표 쪽은 그대로
    finally:
        doc2.close()


def test_docx_replace_whole_word(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("홍길동님께서 홍길동 씨를 만났다")
    doc.save(src)

    d = load(src)
    try:
        report = d.replace_text("홍길동", "유지수", whole_word=True)
        d.save()
    finally:
        d.close()
    assert report["count"] == 1

    assert Document(str(src)).paragraphs[0].text == \
        "홍길동님께서 유지수 씨를 만났다"


def test_docx_replace_not_found_and_invalid_occurrence(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        r1 = doc.replace_text("존재하지않는텍스트", "x")
        r2 = doc.replace_text("담당자", "x", occurrences=[7])
    finally:
        doc.close()
    assert r1["count"] == 0 and r1["not_found"] is True
    assert r2["count"] == 0
    assert r2["invalid_occurrences"] == [7]
    assert r2["total_occurrences"] == 2


# ==================== DOCX: insert_text ====================

def test_docx_insert_after_anchor(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        # 결재란 본문의 담당자(occurrence 1) 뒤에만 삽입
        report = doc.insert_text(
            "담당자", "유지수", position="after",
            occurrences=[1], separator=" ",
        )
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 1

    d = Document(str(src))
    body_texts = [p.text for p in d.paragraphs]
    assert any(t.startswith("담당자 유지수") for t in body_texts)
    # 표 셀은 불변
    assert d.tables[0].cell(0, 0).text == "담당자"


def test_docx_insert_before_anchor(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("성명: (인)")
    doc.save(src)

    d = load(src)
    try:
        d.insert_text("(인)", "유지수", position="before", separator=" ")
        d.save()
    finally:
        d.close()
    assert Document(str(src)).paragraphs[0].text == "성명: 유지수 (인)"


def test_docx_insert_inherits_anchor_format(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("성명")
    r.bold = True
    p.add_run("             ")  # 뒤 공백 run (서식 없음)
    doc.save(src)

    d = load(src)
    try:
        d.insert_text("성명", " 유지수", position="after")
        d.save()
    finally:
        d.close()

    para = Document(str(src)).paragraphs[0]
    # 삽입문은 앵커 run 에 합쳐져 bold 상속
    assert para.runs[0].text == "성명 유지수"
    assert para.runs[0].bold is True


# ==================== DOCX: get_text_map ====================

def test_docx_get_text_map(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    _make_docx_body(src)
    doc = load(src)
    try:
        tmap = doc.get_text_map()
        matches = doc.find_text("담당자")
    finally:
        doc.close()

    assert tmap["format"] == "docx"
    paras = tmap["paragraphs"]
    texts = {e["text"] for e in paras}
    assert "1. 신청인 정보" in texts
    assert "4. 결재란" in texts
    # 제목 플래그
    headings = [e for e in paras if e.get("is_heading")]
    assert len(headings) == 2
    # 표 셀 문단 포함 + location 좌표계
    cell_entries = [e for e in paras if e["scope"] == "table"]
    assert any(e["location"].startswith("table[0].cell(0,0)")
               for e in cell_entries)
    # find_text 의 para_index 와 같은 좌표계
    by_idx = {e["para_index"]: e for e in paras}
    for m in matches:
        assert m.para_index in by_idx
        assert "담당자" in by_idx[m.para_index]["text"]


def test_docx_get_text_map_truncate_and_paging(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("가" * 200)
    for i in range(10):
        doc.add_paragraph(f"문단 {i}")
    doc.save(src)

    d = load(src)
    try:
        tmap = d.get_text_map(max_para_len=50, offset=0, limit=5)
    finally:
        d.close()
    assert tmap["returned"] == 5
    assert tmap["listed"] == 11
    first = tmap["paragraphs"][0]
    assert len(first["text"]) == 50
    assert first["truncated"] is True
    assert first["char_count"] == 200


# ==================== DOCX: 머리말 ====================

def test_docx_header_replace(tmp_path: Path) -> None:
    src = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("본문")
    sec = doc.sections[0]
    sec.header.is_linked_to_previous = False
    sec.header.paragraphs[0].text = "회사명: 홍길동상사"
    doc.save(src)

    d = load(src)
    try:
        matches = d.find_text("홍길동상사")
        assert len(matches) == 1
        assert matches[0].scope == "header"
        d.replace_text("홍길동상사", "유지수상사")
        d.save()
    finally:
        d.close()

    d2 = Document(str(src))
    assert d2.sections[0].header.paragraphs[0].text == "회사명: 유지수상사"


# ==================== HWPX ====================

def _make_hwpx_body(path: Path) -> None:
    doc = HwpxDocument.new()
    doc.add_paragraph("1. 신청인 정보")
    doc.add_paragraph("홍길동님께 안내드립니다")
    doc.add_paragraph("")
    doc.add_table(2, 2)
    doc.save_to_path(path)

    doc2 = HwpxDocument.open(path)
    try:
        section = doc2.sections[0]
        tbl = None
        for para in section.paragraphs:
            if para.tables:
                tbl = para.tables[0]
                break
        assert tbl is not None
        tbl.rows[0].cells[0].text = "담당자"
        tbl.rows[0].cells[1].text = "내용"
        tbl.rows[1].cells[0].text = "홍길동"
        tbl.rows[1].cells[1].text = "비고"
        doc2.save_to_path(path)
    finally:
        doc2.close()


def test_hwpx_find_text_scopes_and_locations(tmp_path: Path) -> None:
    src = tmp_path / "t.hwpx"
    _make_hwpx_body(src)
    doc = load(src)
    try:
        matches = doc.find_text("홍길동")
    finally:
        doc.close()
    assert len(matches) == 2
    assert matches[0].scope == "body"
    assert matches[1].scope == "table"
    assert matches[1].location.startswith("table[0].cell(1,0)")
    # HWPX 는 heading 없음 → context_before 로 섹션 판단
    assert matches[0].context_before == "1. 신청인 정보"


def test_hwpx_replace_persists_after_save(tmp_path: Path) -> None:
    """mark_dirty 누락 시 저장이 안 되는 회귀를 잡는 핵심 테스트."""
    src = tmp_path / "t.hwpx"
    _make_hwpx_body(src)

    doc = load(src)
    try:
        report = doc.replace_text("홍길동", "유지수")
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 2

    doc2 = load(src)
    try:
        assert doc2.find_text("홍길동") == []
        found = doc2.find_text("유지수")
        assert len(found) == 2
        assert doc2.get_cell(0, 1, 0).text == "유지수"
    finally:
        doc2.close()


def test_hwpx_replace_run_split(tmp_path: Path) -> None:
    """<hp:t> 가 쪼개진 문단에서도 concat 기준 매치 + run 보존."""
    src = tmp_path / "t.hwpx"
    _make_hwpx_body(src)

    doc = load(src)
    try:
        # 본문 '홍길동님께 안내드립니다' 문단의 t 를 '홍길'/'동님께 ...' 로 분할
        from lxml import etree
        split_done = False
        for _sec, root in doc._pkg.iter_section_roots():
            for p in root.iter():
                if not p.tag.endswith("}p"):
                    continue
                for run in p.findall(HP_RUN):
                    for t in run.findall(HP_T):
                        if t.text and t.text.startswith("홍길동님께"):
                            rest = t.text[2:]
                            t.text = t.text[:2]
                            t2 = etree.SubElement(run, HP_T)
                            t2.text = rest
                            split_done = True
        assert split_done

        report = doc.replace_text("홍길동님", "유지수님")
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 1

    doc2 = load(src)
    try:
        tmap = doc2.get_text_map()
        texts = [e["text"] for e in tmap["paragraphs"]]
        assert "유지수님께 안내드립니다" in texts
    finally:
        doc2.close()


def test_hwpx_insert_after_in_table_cell(tmp_path: Path) -> None:
    src = tmp_path / "t.hwpx"
    _make_hwpx_body(src)
    doc = load(src)
    try:
        report = doc.insert_text(
            "담당자", "유지수", position="after",
            separator=" ", scope="table",
        )
        doc.save()
    finally:
        doc.close()
    assert report["count"] == 1

    doc2 = load(src)
    try:
        assert doc2.get_cell(0, 0, 0).text == "담당자 유지수"
    finally:
        doc2.close()


def test_hwpx_get_text_map(tmp_path: Path) -> None:
    src = tmp_path / "t.hwpx"
    _make_hwpx_body(src)
    doc = load(src)
    try:
        tmap = doc.get_text_map()
    finally:
        doc.close()
    assert tmap["format"] == "hwpx"
    texts = [e["text"] for e in tmap["paragraphs"]]
    assert "1. 신청인 정보" in texts
    assert "홍길동님께 안내드립니다" in texts
    scopes = {e["scope"] for e in tmap["paragraphs"]}
    assert "table" in scopes and "body" in scopes


# ==================== call_tool dispatcher ====================

def test_call_tool_text_ops_roundtrip(tmp_path: Path) -> None:
    from document_adapter.tools import call_tool

    src = tmp_path / "t.docx"
    _make_docx_body(src)

    tmap = call_tool("get_text_map", {"path": str(src)})
    assert tmap["format"] == "docx" and tmap["returned"] > 0

    found = call_tool("find_text", {"path": str(src), "query": "담당자"})
    assert found["match_count"] == 2

    out = tmp_path / "out.docx"
    rep = call_tool("replace_text", {
        "path": str(src), "old": "홍길동", "new": "유지수",
        "output_path": str(out),
    })
    assert rep["count"] == 1 and rep["output_path"] == str(out)
    # 원본 불변, output 에만 반영
    assert call_tool("find_text",
                     {"path": str(src), "query": "홍길동"})["match_count"] == 1
    assert call_tool("find_text",
                     {"path": str(out), "query": "유지수"})["match_count"] == 1

    ins = call_tool("insert_text", {
        "path": str(out), "anchor": "유지수", "text": "(대리)",
        "separator": " ", "occurrences": [0],
    })
    assert ins["count"] == 1
    assert call_tool("find_text",
                     {"path": str(out),
                      "query": "유지수 (대리)"})["match_count"] == 1


def test_call_tool_text_ops_error_serialization(tmp_path: Path) -> None:
    """미지원 포맷/빈 쿼리도 예외가 아닌 dict 로 직렬화되는지."""
    from pptx import Presentation
    from document_adapter.tools import call_tool

    src = tmp_path / "t.pptx"
    Presentation().save(str(src))
    r = call_tool("find_text", {"path": str(src), "query": "x"})
    assert r["error"] == "not_implemented"

    src2 = tmp_path / "t.docx"
    _make_docx_body(src2)
    r2 = call_tool("find_text", {"path": str(src2), "query": ""})
    assert r2["error"] == "ValueError"


# ==================== 미지원 포맷 ====================

@pytest.mark.parametrize("op", ["find", "replace", "insert", "map"])
def test_pptx_xlsx_not_implemented(tmp_path: Path, op: str) -> None:
    from pptx import Presentation
    src = tmp_path / "t.pptx"
    Presentation().save(str(src))

    doc = load(src)
    try:
        with pytest.raises(NotImplementedForFormat):
            if op == "find":
                doc.find_text("x")
            elif op == "replace":
                doc.replace_text("x", "y")
            elif op == "insert":
                doc.insert_text("x", "y")
            else:
                doc.get_text_map()
    finally:
        doc.close()
