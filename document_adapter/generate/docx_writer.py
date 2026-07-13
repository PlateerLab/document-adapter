"""블록 IR → DOCX 결정적 렌더러 (LLM 없음).

edit2docs v0.9 `documents/docx_engine.docx_from_markdown`(Apache-2.0)의
렌더 전략을 블록 IR 기반으로 이식 — NOTICE 참조.

원칙:
- python-docx **기본 템플릿의 내장 스타일만** 사용 (Heading N / List Bullet /
  List Number / Intense Quote / Table Grid). 커스텀 스타일 정의 금지 —
  로케일별 스타일명 이슈를 차단하고 Word 가 항상 열 수 있는 문서를 보장.
- CJK 로케일은 `w:eastAsia` rFonts 를 Normal 스타일에 명시해야 한글이
  기본 폰트로 렌더된다 (누락 시 Word 가 임의 폰트로 대체).
- 생성 직후 DocxAdapter.load() roundtrip 이 성립해야 한다 (표 좌표 편집).
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt

from .markdown_parser import Block, Span, parse_markdown

__all__ = ["docx_from_markdown", "base_font_for_lang"]

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# 로케일 접두 → 본문 서체. 새 문서에만 적용된다.
_BASE_FONTS = {
    "ko": "맑은 고딕",
    "ja": "Yu Gothic",
    "zh": "Microsoft YaHei",
}
_DEFAULT_BASE_FONT = "Calibri"
_CODE_FONT = "D2Coding"


def base_font_for_lang(lang: str | None) -> str:
    """BCP-47 로케일 → 기본 본문 서체 (ko → 맑은 고딕, 그 외 Calibri)."""
    prefix = (lang or "").split("-")[0].lower()
    return _BASE_FONTS.get(prefix, _DEFAULT_BASE_FONT)


def _add_spans(paragraph, spans: tuple[Span, ...]) -> None:
    for span in spans:
        run = paragraph.add_run(span.text)
        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True
        if span.code:
            run.font.name = _CODE_FONT


def _add_hr(document) -> None:
    """수평선: 빈 문단에 아래 테두리(pBdr/bottom) 주입."""
    p = document.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(f"{{{_W_NS}}}pBdr", {})
    bottom = p_pr.makeelement(
        f"{{{_W_NS}}}bottom",
        {
            f"{{{_W_NS}}}val": "single",
            f"{{{_W_NS}}}sz": "6",
            f"{{{_W_NS}}}space": "1",
            f"{{{_W_NS}}}color": "auto",
        },
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_table(document, block: Block) -> None:
    n_rows = len(block.rows)
    n_cols = len(block.rows[0]) if block.rows else 0
    if n_rows == 0 or n_cols == 0:
        return
    table = document.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    for r, row in enumerate(block.rows):
        for c, cell_spans in enumerate(row):
            cell = table.cell(r, c)
            para = cell.paragraphs[0]
            _add_spans(para, cell_spans)
            if r == 0:  # 헤더행은 굵게 (셀 안 모든 run)
                for run in para.runs:
                    run.bold = True


def docx_from_markdown(
    markdown: str, *, lang: str | None = "ko", base_font: str | None = None
) -> bytes:
    """markdown 서브셋 → 스타일 잡힌 .docx bytes.

    Raises:
        ValueError: 본문이 비어 있을 때 (이중어 메시지 — 호출 레이어의
            재시도 계약이 이 메시지를 LLM 리마인더로 사용한다).
    """
    if not markdown or not markdown.strip():
        raise ValueError(
            "empty document body — provide markdown content. "
            "문서 본문이 비어 있습니다 — markdown 내용을 작성하세요."
        )

    blocks = parse_markdown(markdown)
    if not blocks:
        raise ValueError(
            "markdown produced no renderable blocks. "
            "markdown 에서 렌더 가능한 블록을 찾지 못했습니다."
        )

    font = base_font or base_font_for_lang(lang)
    document = Document()
    style = document.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(10.5)
    # 동아시아 서체는 rPr/rFonts 의 eastAsia 속성에 명시해야 적용된다.
    style.element.rPr.rFonts.set(f"{{{_W_NS}}}eastAsia", font)

    for block in blocks:
        if block.kind == "heading":
            p = document.add_paragraph(style=f"Heading {min(block.level, 6)}")
            _add_spans(p, block.spans)
        elif block.kind == "bullet":
            p = document.add_paragraph(style="List Bullet")
            _add_spans(p, block.spans)
        elif block.kind == "numbered":
            p = document.add_paragraph(style="List Number")
            _add_spans(p, block.spans)
        elif block.kind == "quote":
            p = document.add_paragraph(style="Intense Quote")
            _add_spans(p, block.spans)
        elif block.kind == "hr":
            _add_hr(document)
        elif block.kind == "code":
            for line in block.lines:
                p = document.add_paragraph()
                run = p.add_run(line)
                run.font.name = _CODE_FONT
                run.font.size = Pt(9)
        elif block.kind == "table":
            _add_table(document, block)
        else:  # paragraph (+ 알려지지 않은 kind 안전망)
            p = document.add_paragraph()
            _add_spans(p, block.spans)

    import io

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
