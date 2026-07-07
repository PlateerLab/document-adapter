"""DOCX 어댑터: python-docx (편집) + docxtpl (템플릿 렌더).

표 구조:
- python-docx의 ``row.cells[col]``은 병합된 셀에 대해 동일한 ``_tc``를 여러 번 반환한다.
  이 성질을 이용해 (row, col) → ``_tc`` 매핑을 만든 뒤, 동일 ``_tc``가 등장한
  position들의 bounding box로 병합 anchor/span을 계산한다.
- 중첩 테이블은 ``cell.tables``를 DFS로 훑어 flat index를 부여.
"""
from __future__ import annotations

import re
import warnings
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docxtpl import DocxTemplate

from .base import (
    CellContent,
    CellOutOfBoundsError,
    DocumentAdapter,
    MergeInfo,
    MergedCellWriteError,
    NotImplementedForFormat,
    TableIndexError,
    TableSchema,
    _ParaHandle,
)

TAG_PATTERN = re.compile(r"\{\{\s*(\w+)\s*\}\}")
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# OOXML EMU → cm (1 cm = 360000 EMU)
_EMU_PER_CM = 360000


def _emu_to_cm(emu: Any) -> float | None:
    if emu is None:
        return None
    try:
        v = int(emu)
    except (TypeError, ValueError):
        return None
    if v <= 0:
        return None
    return round(v / _EMU_PER_CM, 1)


def _build_grid(table) -> tuple[dict[tuple[int, int], dict], int, int]:
    """(row,col) → {anchor, span, is_anchor, cell} 매핑.

    python-docx 의 ``row.cells`` 는 가로(gridSpan)+세로(vMerge) 병합이 섞인 표에서
    내부적으로 깨진다(``tc_at_grid_offset`` ValueError). 그래서 OOXML 레이어
    (tblGrid/tr/tc + gridSpan + vMerge)에서 그리드를 직접 계산한다 — hwpx_core
    의 grid 처리와 동일한 접근. 각 행의 tc gridSpan 합이 곧 전체 열 수다.
    """
    from docx.oxml.ns import qn
    from docx.table import _Cell

    tbl = table._tbl
    grid_el = tbl.tblGrid
    n_cols = len(grid_el.gridCol_lst) if grid_el is not None else 0
    trs = tbl.tr_lst
    n_rows = len(trs)
    if n_rows == 0 or n_cols == 0:
        return {}, n_rows, n_cols

    # (r,c) → (anchor_info, is_anchor). anchor_info 는 병합 셀들이 공유하는
    # dict 로, rowspan 이 자라면 span 을 in-place 로 갱신한다.
    cells_info: dict[tuple[int, int], tuple[dict, bool]] = {}
    vmerge_active: dict[int, dict] = {}   # 시작열 → 진행 중인 세로병합 anchor_info

    for r, tr in enumerate(trs):
        col = 0
        for tc in tr.tc_lst:
            gs = tc.grid_span or 1
            tc_pr = tc.tcPr
            vmerge = tc_pr.find(qn("w:vMerge")) if tc_pr is not None else None
            vval = vmerge.get(qn("w:val")) if vmerge is not None else None
            is_continue = vmerge is not None and vval != "restart"

            if is_continue and col in vmerge_active:
                ai = vmerge_active[col]
                ai["span"] = (r - ai["anchor"][0] + 1, ai["span"][1])
                for cc in range(col, col + gs):
                    cells_info[(r, cc)] = (ai, False)
            else:
                ai = {"anchor": (r, col), "span": (1, gs), "cell": _Cell(tc, table)}
                for cc in range(col, col + gs):
                    cells_info[(r, cc)] = (ai, cc == col)
                if vval == "restart":
                    vmerge_active[col] = ai
                else:
                    vmerge_active.pop(col, None)
            col += gs

    grid: dict[tuple[int, int], dict] = {}
    for (r, c), (ai, is_anchor) in cells_info.items():
        grid[(r, c)] = {
            "anchor": ai["anchor"],
            "span": ai["span"],
            "is_anchor": is_anchor,
            "cell": ai["cell"],
        }
    return grid, n_rows, n_cols


class DocxAdapter(DocumentAdapter):
    format = "docx"

    def _open(self) -> None:
        self._doc = Document(str(self.path))

    def save(self, path: Path | str | None = None) -> Path:
        target = Path(path) if path else self.path
        self._doc.save(str(target))
        self.path = target
        return target

    # ---- helpers ----
    def _iter_tables(self) -> Iterator[tuple[int, Any, str]]:
        """Flat DFS (outer + nested). 각 yield: (flat_index, table, parent_path)."""
        idx_counter = [0]

        def walk(tbl, parent_path: str) -> Iterator[tuple[int, Any, str]]:
            current_idx = idx_counter[0]
            idx_counter[0] += 1
            yield current_idx, tbl, parent_path

            grid, _, _ = _build_grid(tbl)
            seen_tc: set[int] = set()
            for (r, c), info in grid.items():
                if not info["is_anchor"]:
                    continue
                tc_key = id(info["cell"]._tc)
                if tc_key in seen_tc:
                    continue
                seen_tc.add(tc_key)
                for nested in info["cell"].tables:
                    child_parent = (
                        f"{parent_path}.tables[{current_idx}].cell({r},{c})"
                    )
                    yield from walk(nested, child_parent)

        for tbl in self._doc.tables:
            yield from walk(tbl, "")

    def _get_table(self, table_index: int):
        for idx, tbl, _ in self._iter_tables():
            if idx == table_index:
                return tbl
        raise TableIndexError(f"DOCX table index {table_index} not found")

    def _resolve_anchor_cell(
        self, tbl, row: int, col: int, *, allow_merge_redirect: bool
    ) -> tuple[Any, dict]:
        """(row,col) → (cell, grid_info). non-anchor 정책 처리."""
        grid, n_rows, n_cols = _build_grid(tbl)
        if row < 0 or col < 0 or row >= n_rows or col >= n_cols:
            raise CellOutOfBoundsError(
                f"cell ({row},{col}) out of bounds ({n_rows}x{n_cols})"
            )
        info = grid.get((row, col))
        if info is None:
            raise CellOutOfBoundsError(
                f"cell ({row},{col}) does not resolve to any physical cell"
            )
        if not info["is_anchor"]:
            anchor_r, anchor_c = info["anchor"]
            if not allow_merge_redirect:
                raise MergedCellWriteError(
                    f"cell ({row},{col}) is part of a merged region anchored at "
                    f"({anchor_r},{anchor_c}) span={info['span']}. "
                    f"Write to the anchor coordinate, or pass "
                    f"allow_merge_redirect=True."
                )
            warnings.warn(
                f"write to ({row},{col}) redirected to merge anchor "
                f"({anchor_r},{anchor_c})",
                stacklevel=3,
            )
        return info["cell"], info

    # ---- inspection ----
    def get_placeholders(self) -> list[str]:
        keys: set[str] = set()
        for p in self._doc.paragraphs:
            keys.update(TAG_PATTERN.findall(p.text))
        # 머리말/꼬리말 (docxtpl render 는 이미 채우므로 inspect 와 일치시킨다)
        for section in self._doc.sections:
            for hf in (section.header, section.footer):
                for p in hf.paragraphs:
                    keys.update(TAG_PATTERN.findall(p.text))
        # 모든 (중첩 포함) 표 셀에서 수집. row.cells 는 병합표에서 깨지므로
        # _build_grid 의 anchor 셀만 순회한다(get_tables 와 동일 견고 경로).
        for _, tbl, _ in self._iter_tables():
            grid, _, _ = _build_grid(tbl)
            seen: set[int] = set()
            for info in grid.values():
                if not info["is_anchor"]:
                    continue
                tc_id = id(info["cell"]._tc)
                if tc_id in seen:
                    continue
                seen.add(tc_id)
                keys.update(TAG_PATTERN.findall(info["cell"].text))
        return sorted(keys)

    def get_tables(self, min_rows: int = 1, min_cols: int = 1,
                   preview_rows: int = 4, max_cell_len: int = 40) -> list[TableSchema]:
        schemas: list[TableSchema] = []
        for idx, tbl, parent_path in self._iter_tables():
            grid, n_rows, n_cols = _build_grid(tbl)
            if n_rows < min_rows or n_cols < min_cols:
                continue

            visible_rows = min(n_rows, preview_rows)
            preview: list[list[str | None]] = [
                [None for _ in range(n_cols)] for _ in range(visible_rows)
            ]
            merges: list[MergeInfo] = []
            seen_anchors: set[tuple[int, int]] = set()

            # grid 순회 — 앵커 위치에만 텍스트 주입
            for (r, c), info in sorted(grid.items()):
                if info["anchor"] in seen_anchors:
                    continue
                if info["is_anchor"]:
                    seen_anchors.add(info["anchor"])
                    if r < visible_rows:
                        text = (info["cell"].text or "").strip()
                        preview[r][c] = text[:max_cell_len]
                    if info["span"] != (1, 1):
                        merges.append(MergeInfo(anchor=info["anchor"], span=info["span"]))

            # 셀 크기 힌트
            col_widths = [
                _emu_to_cm(getattr(col, "width", None)) for col in tbl.columns
            ]
            row_heights = [
                _emu_to_cm(getattr(row, "height", None)) for row in tbl.rows
            ]
            col_widths_out = col_widths if any(v is not None for v in col_widths) else None
            row_heights_out = row_heights if any(v is not None for v in row_heights) else None

            schemas.append(
                TableSchema(
                    index=idx,
                    rows=n_rows,
                    cols=n_cols,
                    preview=preview,
                    merges=merges,
                    parent_path=parent_path or None,
                    column_widths_cm=col_widths_out,
                    row_heights_cm=row_heights_out,
                )
            )
        return schemas

    def get_cell(self, table_index: int, row: int, col: int) -> CellContent:
        tbl = self._get_table(table_index)
        grid, n_rows, n_cols = _build_grid(tbl)
        if row < 0 or col < 0 or row >= n_rows or col >= n_cols:
            raise CellOutOfBoundsError(
                f"cell ({row},{col}) out of bounds ({n_rows}x{n_cols})"
            )
        info = grid.get((row, col))
        if info is None:
            raise CellOutOfBoundsError(
                f"cell ({row},{col}) does not resolve to any physical cell"
            )

        cell = info["cell"]
        paragraphs_text = [p.text for p in cell.paragraphs]
        text = cell.text or ""

        nested_indices: list[int] = []
        if info["is_anchor"] and list(cell.tables):
            nested_tc_ids = {id(t._tbl) for t in cell.tables}
            for child_idx, child_tbl, _ in self._iter_tables():
                if id(child_tbl._tbl) in nested_tc_ids:
                    nested_indices.append(child_idx)

        # 셀 크기 힌트 (anchor 기준 span 영역 합)
        a_r, a_c = info["anchor"]
        r_span, c_span = info["span"]
        try:
            cols_list = list(tbl.columns)
            rows_list = list(tbl.rows)
            width_emu = sum(
                getattr(cols_list[i], "width", 0) or 0
                for i in range(a_c, min(a_c + c_span, len(cols_list)))
            )
            height_emu = sum(
                getattr(rows_list[i], "height", 0) or 0
                for i in range(a_r, min(a_r + r_span, len(rows_list)))
            )
        except (IndexError, AttributeError):
            width_emu = 0
            height_emu = 0

        return CellContent(
            row=row,
            col=col,
            text=text,
            paragraphs=paragraphs_text,
            is_anchor=info["is_anchor"],
            anchor=info["anchor"],
            span=info["span"],
            nested_table_indices=nested_indices,
            width_cm=_emu_to_cm(width_emu),
            height_cm=_emu_to_cm(height_emu),
            char_count=len(text),
        )

    # ---- editing ----
    def render_template(self, context: dict[str, Any], *,
                        on_missing: str = "blank") -> dict[str, list[str]]:
        """docxtpl 기반 Jinja2 렌더. 누락 키는 on_missing 정책(base 참조).
        - `{%tr for row in rows %}` / `{%tr endfor %}`는 **각각 별도 행**에 두어야 함
        - 같은 행에 두면 `<w:tr>` 전체가 `{% for %}`로 교체되어 endfor 손실
        """
        report = self._render_report(self.get_placeholders(), context, on_missing)
        tpl = DocxTemplate(self.path)
        if on_missing == "leave":
            import jinja2
            env = jinja2.Environment(undefined=jinja2.DebugUndefined,
                                     autoescape=True)
            tpl.render(context, jinja_env=env)
        else:
            # blank: docxtpl 기본(Jinja Undefined→""). error: 위에서 이미 raise.
            tpl.render(context)
        tpl.save(self.path)
        self._doc = Document(str(self.path))
        return report

    def set_cell(
        self,
        table_index: int,
        row: int,
        col: int,
        value: str,
        *,
        allow_merge_redirect: bool = False,
    ) -> str:
        tbl = self._get_table(table_index)
        cell, info = self._resolve_anchor_cell(
            tbl, row, col, allow_merge_redirect=allow_merge_redirect
        )
        old = cell.text
        _set_cell_preserving_format(cell, value)
        return old

    def append_to_cell(
        self,
        table_index: int,
        row: int,
        col: int,
        value: str,
        separator: str = "  ",
        *,
        allow_merge_redirect: bool = False,
    ) -> str:
        tbl = self._get_table(table_index)
        cell, info = self._resolve_anchor_cell(
            tbl, row, col, allow_merge_redirect=allow_merge_redirect
        )
        old = cell.text
        new_value = f"{old}{separator}{value}" if old else value
        _set_cell_preserving_format(cell, new_value)
        return old

    def append_row(self, table_index: int, values: list[str]) -> None:
        # v0.14: python-docx add_row() 는 서식 없는 빈 행을 만들었다 (음영·
        # 테두리·행높이·폰트 미상속). 마지막 행 deepcopy 방식의 insert_row 로
        # 위임해 HWPX 와 동작을 통일한다.
        self.insert_row(table_index, values, at_row=None)

    # ---- 위치 지정 행/열 삽입 (v0.14+) ----

    def insert_row(
        self,
        table_index: int,
        values: list[str],
        at_row: int | None = None,
    ) -> int:
        """지정 위치에 새 행 삽입 — 인접 행 <w:tr> deepcopy 로 서식 상속.

        - 템플릿 행: at_row 위치의 기존 행 (맨 끝 삽입이면 마지막 행).
          trPr(행높이)·tcPr(음영/테두리/폭)·pPr/rPr(폰트) 전부 상속된다.
        - 복사된 셀의 vMerge 는 제거 (새 행은 항상 독립 행), 중첩 표 제거.
        - 삽입 경계를 세로 병합이 가로지르면 NotImplementedForFormat.
        - values 는 새 행의 **물리적 셀** 순서 (gridSpan 병합 셀은 1개로 침).
        """
        tbl = self._get_table(table_index)
        grid, n_rows, n_cols = _build_grid(tbl)
        if n_rows == 0:
            raise NotImplementedForFormat(
                "cannot insert a row into an empty DOCX table"
            )
        if at_row is None:
            at_row = n_rows
        if at_row < 0 or at_row > n_rows:
            raise CellOutOfBoundsError(
                f"at_row {at_row} out of range 0..{n_rows}"
            )

        # 세로 병합 경계 가드: 삽입 지점 위아래가 같은 병합 영역이면 거부
        if 0 < at_row < n_rows:
            for c in range(n_cols):
                info = grid.get((at_row, c))
                if info is not None and info["anchor"][0] < at_row:
                    raise NotImplementedForFormat(
                        f"insertion point row {at_row} crosses a vertical "
                        f"merge anchored at {info['anchor']}; inserting here "
                        f"would split the merged region."
                    )

        trs = tbl._tbl.tr_lst
        template_idx = at_row if at_row < n_rows else n_rows - 1
        new_tr = deepcopy(trs[template_idx])
        for tc in new_tr.tc_lst:
            _blank_copied_tc(tc)
            _strip_tc_props(tc, "w:vMerge")

        if at_row < n_rows:
            trs[at_row].addprevious(new_tr)
        else:
            trs[-1].addnext(new_tr)

        # 값 채우기 — 논리 grid 열 기준 (HWPX 와 동일 규약).
        # gridSpan 병합 셀은 anchor 열의 값을 받고, 병합에 덮인 열 값은 무시.
        from docx.table import _Cell
        col = 0
        for tc in new_tr.tc_lst:
            gs = tc.grid_span or 1
            if col < len(values) and values[col]:
                _set_cell_preserving_format(_Cell(tc, tbl), values[col])
            col += gs
        return at_row

    def insert_column(
        self,
        table_index: int,
        values: list[str],
        at_col: int | None = None,
    ) -> int:
        """지정 위치에 새 열 삽입 — 행별 인접 셀 <w:tc> deepcopy 로 서식 상속.

        - 템플릿 셀: 각 행에서 왼쪽 이웃 열의 셀 (at_col=0 이면 오른쪽 이웃).
          행별로 복사하므로 헤더 행은 헤더 서식, 데이터 행은 데이터 서식.
        - 복사된 셀의 gridSpan/vMerge 제거 (새 열은 1칸짜리 독립 셀), 중첩 표 제거.
        - 표 전체 폭 유지: tblGrid/tcW 폭을 비례 축소해 새 열 폭을 흡수.
        - 삽입 경계를 가로 병합(gridSpan)이 가로지르는 행이 있으면 거부.
        - values 는 위 행부터 (values[0] 이 보통 헤더).
        """
        from docx.oxml.ns import qn
        from docx.table import _Cell

        tbl = self._get_table(table_index)
        grid, n_rows, n_cols = _build_grid(tbl)
        if n_rows == 0 or n_cols == 0:
            raise NotImplementedForFormat(
                "cannot insert a column into an empty DOCX table"
            )
        if at_col is None:
            at_col = n_cols
        if at_col < 0 or at_col > n_cols:
            raise CellOutOfBoundsError(
                f"at_col {at_col} out of range 0..{n_cols}"
            )

        # 가로 병합 경계 가드
        if 0 < at_col < n_cols:
            for r in range(n_rows):
                info = grid.get((r, at_col))
                if info is not None and info["anchor"][1] < at_col:
                    raise NotImplementedForFormat(
                        f"insertion point column {at_col} crosses a horizontal "
                        f"merge anchored at {info['anchor']} in row {r}; "
                        f"inserting here would split the merged region."
                    )

        template_col = at_col - 1 if at_col > 0 else 0

        # ── tblGrid 갱신 + 폭 비례 재배분 (표 전체 폭 유지) ──
        grid_el = tbl._tbl.tblGrid
        grid_cols = grid_el.gridCol_lst
        widths: list[int | None] = []
        for gc in grid_cols:
            w = gc.get(qn("w:w"))
            try:
                widths.append(int(w) if w is not None else None)
            except (TypeError, ValueError):
                widths.append(None)

        new_gc = deepcopy(grid_cols[template_col])
        scale = 1.0
        if all(w is not None and w > 0 for w in widths):
            total = sum(widths)  # type: ignore[arg-type]
            new_w = widths[template_col]
            scale = total / (total + new_w)  # type: ignore[operator]
            for gc, w in zip(grid_cols, widths):
                gc.set(qn("w:w"), str(int(round(w * scale))))  # type: ignore[operator]
            new_gc.set(qn("w:w"), str(int(round(new_w * scale))))  # type: ignore[operator]
        if at_col < n_cols:
            grid_cols[at_col].addprevious(new_gc)
        else:
            grid_cols[-1].addnext(new_gc)

        # ── 행별 셀 삽입 ──
        def _scale_tc_width(tc) -> None:
            if scale == 1.0:
                return
            tc_pr = tc.tcPr
            tcw = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
            if tcw is not None and tcw.get(qn("w:type")) == "dxa":
                w = tcw.get(qn("w:w"))
                try:
                    tcw.set(qn("w:w"), str(int(round(int(w) * scale))))
                except (TypeError, ValueError):
                    pass

        new_tcs: list[Any] = []
        for tr in tbl._tbl.tr_lst:
            tcs = tr.tc_lst
            if not tcs:
                new_tcs.append(None)
                continue
            # 물리 셀 순회로 grid 좌표 계산: 삽입 대상/템플릿 셀 탐색
            col = 0
            insert_before = None
            template_tc = tcs[0]
            for tc in tcs:
                gs = tc.grid_span or 1
                if col <= template_col < col + gs:
                    template_tc = tc
                if col >= at_col and insert_before is None:
                    insert_before = tc
                col += gs
                _scale_tc_width(tc)

            new_tc = deepcopy(template_tc)
            _blank_copied_tc(new_tc)
            _strip_tc_props(new_tc, "w:vMerge", "w:gridSpan")
            _scale_tc_width(new_tc)
            if insert_before is not None:
                insert_before.addprevious(new_tc)
            else:
                tcs[-1].addnext(new_tc)
            new_tcs.append(new_tc)

        for r, tc in enumerate(new_tcs):
            if tc is not None and r < len(values) and values[r]:
                _set_cell_preserving_format(_Cell(tc, tbl), values[r])
        return at_col

    # ---- paragraph text ops (v0.13+) ----

    def _iter_text_paragraphs(self):
        """본문(표와 문서 순서 교차)+표 셀+머리말/꼬리말 문단을 순회.

        - 본문/표 순서: ``iter_inner_content`` (python-docx 1.1+) 로 문서
          등장 순서를 보존 — nearest_heading 이 표 안 매치에도 올바르게
          연결된다.
        - 표 위치 표기는 set_cell 과 동일한 flat table_index 좌표계.
          ``_iter_tables`` 는 DFS 라 top-level 표의 중첩 subtree 가 연속
          블록으로 나오는 성질을 이용해 매핑한다.
        - run 열거: 직접 run + hyperlink 내부 run (문서 순서 유지).
          ``.//w:r`` 을 쓰지 않는 이유 — run 안에 중첩된 텍스트박스
          (``w:txbxContent``)의 다른 문단 run 까지 끌려 들어오기 때문.
        - 머리말/꼬리말: linked(자체 정의 없음) 는 건너뜀 — 이전 섹션
          내용의 중복이자, 접근 시 정의가 생성될 수 있어 문서를 오염시킨다.
        """
        from docx.table import Table as _Table
        from docx.text.paragraph import Paragraph as _Paragraph
        from docx.text.run import Run as _Run

        def para_handle(para, scope: str, location: str) -> _ParaHandle:
            r_els = para._p.xpath("./w:r | ./w:hyperlink/w:r")
            runs = [_Run(r, para) for r in r_els]

            def get_texts() -> list[str]:
                return [r.text for r in runs]

            def set_texts(new_texts: list[str]) -> None:
                # 값이 달라진 run 만 재기록 — 무관한 run 의 XML 재구성
                # (탭/개행 요소 재생성 등)을 피한다.
                for r, new_t in zip(runs, new_texts):
                    if r.text != new_t:
                        r.text = new_t

            return _ParaHandle(
                scope=scope,
                location=location,
                is_heading=_is_heading_para(para),
                get_texts=get_texts,
                set_texts=set_texts,
            )

        # 표 flat index 매핑: top-level 표 → [(flat_idx, tbl), ...(중첩 포함)]
        subtree_by_toplevel: dict[int, list[tuple[int, Any]]] = {}
        current_key: int | None = None
        for idx, tbl, parent_path in self._iter_tables():
            if parent_path == "":
                current_key = id(tbl._tbl)
                subtree_by_toplevel[current_key] = []
            if current_key is not None:
                subtree_by_toplevel[current_key].append((idx, tbl))

        body_para_n = 0
        for item in self._doc.iter_inner_content():
            if isinstance(item, _Paragraph):
                yield para_handle(item, "body", f"body.p[{body_para_n}]")
                body_para_n += 1
            elif isinstance(item, _Table):
                for idx, tbl in subtree_by_toplevel.get(id(item._tbl), []):
                    grid, _, _ = _build_grid(tbl)
                    seen_tc: set[int] = set()
                    for (r, c), info in sorted(grid.items()):
                        if not info["is_anchor"]:
                            continue
                        tc_key = id(info["cell"]._tc)
                        if tc_key in seen_tc:
                            continue
                        seen_tc.add(tc_key)
                        for k, cp in enumerate(info["cell"].paragraphs):
                            yield para_handle(
                                cp, "table",
                                f"table[{idx}].cell({r},{c}).p[{k}]",
                            )

        for si, section in enumerate(self._doc.sections):
            for kind, part in (("header", section.header),
                               ("footer", section.footer)):
                if part.is_linked_to_previous:
                    continue
                for k, hp in enumerate(part.paragraphs):
                    yield para_handle(
                        hp, kind, f"section[{si}].{kind}.p[{k}]",
                    )


def _is_heading_para(para) -> bool:
    """문단이 제목(heading)인지 휴리스틱 판정.

    1. 스타일 이름이 "Heading ..." / "제목 ..." 계열
    2. 문단 자체에 outlineLvl 이 지정된 경우 (스타일 무관 개요 수준)
    """
    try:
        name = (para.style.name or "").lower()
    except Exception:
        name = ""
    if name.startswith("heading") or name.startswith("제목"):
        return True
    from docx.oxml.ns import qn
    ppr = para._p.pPr
    return ppr is not None and ppr.find(qn("w:outlineLvl")) is not None


def _blank_copied_tc(tc) -> None:
    """deepcopy 된 ``<w:tc>`` 를 '빈 셀'로 정리 (서식은 유지).

    - 중첩 표 제거: 템플릿 셀 안의 중첩 표가 새 셀로 복제되는 것 방지.
    - 모든 run 텍스트 비움: ``<w:rPr>`` 는 남으므로 폰트/크기/볼드가 상속되고,
      ``_set_cell_preserving_format`` 이 첫 run 을 재사용해 값을 쓴다.
    """
    from docx.oxml.ns import qn

    for nested in tc.findall(qn("w:tbl")):
        tc.remove(nested)
    for t in tc.iter(qn("w:t")):
        t.text = ""


def _strip_tc_props(tc, *prop_tags: str) -> None:
    """``<w:tcPr>`` 에서 지정 속성 요소 제거 (예: "w:vMerge", "w:gridSpan").

    복사로 만든 새 셀은 병합에 참여하지 않는 독립 셀이어야 한다.
    """
    from docx.oxml.ns import qn

    tc_pr = tc.tcPr
    if tc_pr is None:
        return
    for tag in prop_tags:
        el = tc_pr.find(qn(tag))
        if el is not None:
            tc_pr.remove(el)


def _set_cell_preserving_format(cell, value: str) -> None:
    """Write ``value`` into ``cell`` without dropping run formatting.

    ``python-docx``'s ``cell.text = value`` setter wipes every paragraph and
    run in the cell, replacing them with a brand-new default-styled run. That
    destroys two kinds of formatting:

    1. **Existing runs** — font, size, bold, color on already-populated cells.
    2. **Paragraph mark run properties** — an empty cell often holds a
       ``<w:p><w:pPr><w:rPr>…</w:rPr></w:pPr></w:p>`` describing how the
       next typed character should look. Real templates put the table font
       here so the cell renders correctly even before any text exists.

    Strategy:

    - If any paragraph already has runs, reuse the first one and blank the
      rest.
    - Otherwise, append a new ``<w:r>`` into the first paragraph and clone
      its ``<w:pPr><w:rPr>`` into the new run's ``<w:rPr>`` so the empty-cell
      font survives.

    Paragraph identity is compared by index because python-docx returns a
    fresh Python wrapper on repeated ``cell.paragraphs`` accesses.
    """
    paragraphs = list(cell.paragraphs)
    first_idx = next((i for i, para in enumerate(paragraphs) if para.runs), None)

    if first_idx is not None:
        first_para = paragraphs[first_idx]
        first_para.runs[0].text = value
        for run in first_para.runs[1:]:
            run.text = ""
        for i, para in enumerate(paragraphs):
            if i == first_idx:
                continue
            for run in para.runs:
                run.text = ""
        return

    target_para = paragraphs[0] if paragraphs else None
    if target_para is None:
        cell.text = value
        return

    run = target_para.add_run(value)
    p_el = target_para._p
    ppr = p_el.find(f"{{{_W_NS}}}pPr")
    if ppr is not None:
        rpr_in_ppr = ppr.find(f"{{{_W_NS}}}rPr")
        if rpr_in_ppr is not None:
            cloned = deepcopy(rpr_in_ppr)
            cloned.tag = f"{{{_W_NS}}}rPr"
            run._r.insert(0, cloned)
