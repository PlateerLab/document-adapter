"""DOCX 양식 fixture 합성 — fill_form 시나리오용.

로컬 DOCX 는 전부 기술 문서라 양식성 fixture 가 없음. docx-level label|value
표를 가진 합성 양식을 생성해 시나리오 실험에 사용.
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Cm, Pt


def make_employee_form(dst: Path) -> None:
    """직원 정보 양식 — 라벨|값 4x2 표 + 체크리스트 2개 표."""
    doc = Document()
    doc.add_heading("직원 정보 등록 양식", level=1)
    doc.add_paragraph("아래 양식을 작성해주세요.")

    # 표 1: 기본 정보 (4x2 라벨|값)
    doc.add_paragraph()
    doc.add_heading("1. 기본 정보", level=2)
    t1 = doc.add_table(rows=4, cols=2)
    t1.style = "Table Grid"
    labels_values = [
        ("성명", ""),
        ("부서", ""),
        ("직급", ""),
        ("입사일", ""),
    ]
    for i, (label, val) in enumerate(labels_values):
        t1.cell(i, 0).text = label
        t1.cell(i, 1).text = val
    # column width 설정
    for col in t1.columns:
        col.cells[0].width = Cm(3)

    # 표 2: 연락처 (3x2)
    doc.add_paragraph()
    doc.add_heading("2. 연락처", level=2)
    t2 = doc.add_table(rows=3, cols=2)
    t2.style = "Table Grid"
    for i, (label, val) in enumerate([("전화번호", ""), ("이메일", ""), ("주소", "")]):
        t2.cell(i, 0).text = label
        t2.cell(i, 1).text = val

    # 표 3: 평가 항목 (5x3) — "항목|평가|비고" 형식. fill_form 비적합, set_cell 적합.
    doc.add_paragraph()
    doc.add_heading("3. 평가", level=2)
    t3 = doc.add_table(rows=5, cols=3)
    t3.style = "Table Grid"
    for i, hdr in enumerate(["항목", "평가", "비고"]):
        t3.cell(0, i).text = hdr
    for i, item in enumerate(["업무능력", "협업", "리더십", "성장가능성"], start=1):
        t3.cell(i, 0).text = item

    doc.save(dst)
    print(f"생성: {dst} ({dst.stat().st_size:,} bytes)")


if __name__ == "__main__":
    out_dir = ROOT / "tests" / "fixtures" / "docx" / "real"
    out_dir.mkdir(parents=True, exist_ok=True)
    make_employee_form(out_dir / "employee_form.docx")
