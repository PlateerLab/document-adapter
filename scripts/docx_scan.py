"""로컬 DOCX 를 스캔해 fixture 후보 선별 (표/병합 셀/크기 기준)."""
from __future__ import annotations

import sys
from pathlib import Path

SEARCH_ROOTS = [
    Path.home() / "Documents",
    Path.home() / "Desktop",
    Path.home() / "Downloads",
]


def find_docx() -> list[Path]:
    results = []
    for root in SEARCH_ROOTS:
        if not root.exists():
            continue
        for p in root.rglob("*.docx"):
            if p.name.startswith("~$"):
                continue
            if "node_modules" in p.parts:
                continue
            results.append(p)
    return results


def analyze(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}

    try:
        doc = Document(str(path))
        table_count = len(doc.tables)
        merge_count = 0
        max_rows = 0
        max_cols = 0
        total_cells = 0
        paragraphs = len(doc.paragraphs)

        for table in doc.tables:
            rows = len(table.rows)
            cols = max((len(r.cells) for r in table.rows), default=0)
            max_rows = max(max_rows, rows)
            max_cols = max(max_cols, cols)
            # 병합 감지: 같은 _tc 가 여러 position 에 등장
            seen_tcs = {}
            for r_idx, row in enumerate(table.rows):
                for c_idx in range(min(len(row.cells), cols)):
                    tc = id(row.cells[c_idx]._tc)
                    total_cells += 1
                    if tc in seen_tcs:
                        seen_tcs[tc] += 1
                    else:
                        seen_tcs[tc] = 1
            merge_count += sum(1 for cnt in seen_tcs.values() if cnt > 1)

        return {
            "table_count": table_count,
            "merge_count": merge_count,
            "max_rows": max_rows,
            "max_cols": max_cols,
            "total_cells": total_cells,
            "paragraphs": paragraphs,
            "error": None,
        }
    except Exception as e:
        return {"error": f"{type(e).__name__}: {e}"}


def main() -> int:
    paths = find_docx()
    print(f"총 {len(paths)}개 DOCX 발견, 분석 중...\n")

    rows = []
    for p in paths:
        size = p.stat().st_size
        m = analyze(p)
        rows.append((p, size, m))

    def sort_key(r):
        p, size, m = r
        if m.get("error"):
            return (999, 0, 0)
        # 병합 셀 많은 것 우선, 표 많은 것 우선, 크기 다양
        return (-m.get("merge_count", 0), -m.get("table_count", 0), size)

    rows.sort(key=sort_key)

    print(f"{'Size':>8} {'Tables':>6} {'Merges':>6} {'MaxRC':>8} {'Paras':>6}  Path")
    print("-" * 120)
    for p, size, m in rows[:30]:
        if m.get("error"):
            continue
        size_str = f"{size//1024}K" if size < 1024*1024 else f"{size/(1024*1024):.1f}M"
        maxrc = f"{m['max_rows']}x{m['max_cols']}"
        print(
            f"{size_str:>8} {m['table_count']:>6} {m['merge_count']:>6} "
            f"{maxrc:>8} {m['paragraphs']:>6}  {p}"
        )

    total_with_tables = sum(
        1 for _, _, m in rows if isinstance(m.get("table_count"), int) and m["table_count"] > 0
    )
    total_with_merges = sum(
        1 for _, _, m in rows if isinstance(m.get("merge_count"), int) and m["merge_count"] > 0
    )
    print(f"\n표 있는 파일: {total_with_tables}")
    print(f"병합 셀 있는 파일: {total_with_merges}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
