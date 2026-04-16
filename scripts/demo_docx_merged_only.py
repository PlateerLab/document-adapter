"""x2bee_checklist.docx 의 T0 (13x5 복잡 병합 표) 만 남기고 편집 시연.

T0 구조:
  - 13행 × 5열, 6개 병합:
    * (1,0) span=(6,1) '검색엔진' — 세로 6칸 병합
    * (1,1) span=(5,1) 'Opensearch' — 세로 5칸 병합
    * (7,0) span=(5,1) 'X2BEE' — 세로 5칸 병합
    * (7,1) span=(2,1) 'K8S' — 세로 2칸 병합
    * (9,1) span=(2,1) 'Spring' — 세로 2칸 병합
    * (12,1) span=(1,4) 점검결과 — 가로 4칸 병합
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document

from document_adapter import load

FIXTURE = ROOT / "tests" / "fixtures" / "docx" / "real" / "x2bee_checklist.docx"
OUT_DIR = Path.home() / "Desktop" / "docx_demo"
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_first_table_only(src: Path, dst: Path) -> None:
    """T0 만 남기고 나머지 paragraph / table 전부 삭제."""
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    target_table = doc.tables[0]
    t0_xml = deepcopy(target_table._tbl)

    body = doc.element.body
    sect_pr = None
    for child in list(body):
        if child.tag == f"{W_NS}sectPr":
            sect_pr = child
        body.remove(child)

    doc.add_paragraph("X2BEE 점검 체크리스트 — T0 (13x5, 6 merges)").style = (
        doc.styles["Heading 1"]
    )
    doc.add_paragraph()
    body.append(t0_xml)
    if sect_pr is not None:
        body.append(sect_pr)

    doc.save(str(dst))
    print(f"생성: {dst.name} — T0 표만 보존")


def main() -> int:
    OUT_DIR.mkdir(exist_ok=True)
    # 기존 07, 08 파일 제거
    for p in OUT_DIR.glob("07_*.docx"):
        p.unlink()
    for p in OUT_DIR.glob("08_*.docx"):
        p.unlink()
    for p in OUT_DIR.glob("09_*.docx"):
        p.unlink()
    for p in OUT_DIR.glob("10_*.docx"):
        p.unlink()

    original = OUT_DIR / "09_complex_merged_original.docx"
    extract_first_table_only(FIXTURE, original)

    edited = OUT_DIR / "10_complex_merged_edited.docx"
    shutil.copy2(original, edited)
    a = load(edited)
    try:
        tables = a.get_tables(preview_rows=15, max_cell_len=25)
        t = tables[0]
        print(f"\nT0 구조: {t.rows}x{t.cols}, merges={len(t.merges)}")
        for m in t.merges:
            print(f"  anchor={m.anchor} span={m.span}")

        # 6개 병합 anchor 전부 편집 — 각각 "★" 마커 추가
        edits = [
            ((1, 0), "★ 검색엔진 (세로 6칸) ★"),
            ((1, 1), "★ Opensearch (세로 5칸) ★"),
            ((7, 0), "★ X2BEE (세로 5칸) ★"),
            ((7, 1), "★ K8S (세로 2칸) ★"),
            ((9, 1), "★ Spring (세로 2칸) ★"),
            ((12, 1), "★ 점검결과 라벨 (가로 4칸) ★"),
        ]
        for (r, c), new_val in edits:
            old = a.set_cell(0, r, c, new_val)
            print(f"  ({r},{c}) '{old}' → '{new_val}'")
        a.save(edited)
    finally:
        a.close()

    print(f"\n생성된 파일:")
    for p in sorted(OUT_DIR.glob("0[9]*.docx")) + sorted(OUT_DIR.glob("10_*.docx")):
        print(f"  {p.name} — {p.stat().st_size:,} bytes")
    print(f"\nFinder: open '{OUT_DIR}'")
    return 0


if __name__ == "__main__":
    sys.exit(main())
